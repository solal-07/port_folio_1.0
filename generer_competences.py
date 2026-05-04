"""
Génération des fiches de compétences BTS SIO — Solal BEN ZAQUI
Un document Word par compétence, au format de l'exemple de documentation.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

PYTHON = "/c/Users/Solal/AppData/Local/Programs/Python/Python313/python.exe"
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "fiches_competences")
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ─── Helpers ────────────────────────────────────────────────────────────────

def set_heading_style(paragraph, level=1, color=(30, 30, 30)):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
    run.bold = True
    run.font.size = Pt(18 if level == 1 else 14)
    run.font.color.rgb = RGBColor(*color)


def add_title(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.size = Pt(22)
    p.runs[0].font.color.rgb = RGBColor(31, 78, 121)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_section(doc, title):
    p = doc.add_heading(title, level=2)
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.color.rgb = RGBColor(68, 114, 196)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.runs[0].font.size = Pt(11)
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(40, 40, 40)
    # Light grey background via shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def create_doc():
    doc = Document()
    # Page margins
    section = doc.sections[0]
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)
    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    return doc


def add_header_info(doc, nom_candidat="BEN ZAQUI SOLAL", centre="BTS SIO — EFREI"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Candidat : {nom_candidat}  |  {centre}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)
    doc.add_paragraph()  # espace


# ─── Fiche 1 : Gérer le patrimoine informatique ─────────────────────────────

def fiche_1():
    doc = create_doc()
    add_header_info(doc)

    add_title(doc, "Réalisation professionnelle : Recensement du parc informatique avec GLPI sur VM Debian 12")

    add_section(doc, "Compétences mises en œuvre")
    add_body(doc,
        "• Recenser et identifier les ressources numériques\n"
        "• Exploiter des référentiels, normes et standards adoptés par le prestataire informatique\n"
        "• Mettre en place et vérifier les niveaux d'habilitation associés à un service\n"
        "• Vérifier le respect des règles d'utilisation des ressources numériques"
    )

    add_section(doc, "Contexte")
    add_body(doc,
        "Dans le cadre de ma formation BTS SIO option SISR, j'ai été amené à mettre en place un système "
        "de gestion de parc informatique (GMAO/ITSM) afin d'inventorier et de suivre l'ensemble des "
        "ressources numériques d'une organisation fictive. L'objectif était d'améliorer la traçabilité "
        "des équipements et la gestion des licences logicielles."
    )

    add_section(doc, "Démarche suivie")
    add_body(doc,
        "1. Création et configuration d'une machine virtuelle sous Debian 12 (VirtualBox).\n"
        "2. Installation du serveur LAMP (Apache2, MariaDB, PHP) nécessaire au fonctionnement de GLPI."
    )
    add_code_block(doc,
        "sudo apt update && sudo apt upgrade -y\n"
        "sudo apt install apache2 mariadb-server php php-mysql libapache2-mod-php -y"
    )
    add_body(doc,
        "3. Téléchargement et installation de GLPI 10.x, configuration de la base de données et de l'accès web."
    )
    add_code_block(doc,
        "wget https://github.com/glpi-project/glpi/releases/download/10.0.12/glpi-10.0.12.tgz\n"
        "tar -xzf glpi-10.0.12.tgz -C /var/www/html/\n"
        "sudo chown -R www-data:www-data /var/www/html/glpi"
    )
    add_body(doc,
        "4. Installation et déploiement de l'agent GLPI sur les postes clients pour la remontée automatique "
        "des informations matérielles et logicielles."
    )
    add_code_block(doc,
        "sudo apt install glpi-agent\n"
        "# Configuration de l'URL du serveur GLPI dans /etc/glpi-agent/agent.cfg\n"
        "server = http://192.168.1.100/glpi"
    )
    add_body(doc,
        "5. Création des catégories de matériels (ordinateurs, imprimantes, périphériques réseau), "
        "attribution des habilitations par rôle (technicien, administrateur, utilisateur).\n"
        "6. Vérification de la conformité des licences et application des règles d'utilisation "
        "définies par la charte informatique."
    )

    add_section(doc, "Conclusion")
    add_body(doc,
        "Cette réalisation m'a permis de maîtriser le déploiement complet de GLPI sur un environnement "
        "Linux Debian. J'ai appris à configurer un serveur LAMP, à inventorier automatiquement un parc "
        "via l'agent GLPI, et à gérer les droits d'accès selon les rôles des utilisateurs. "
        "L'outil offre une vision centralisée et en temps réel du patrimoine numérique de l'organisation, "
        "facilitant les décisions de maintenance et de renouvellement du matériel."
    )

    path = os.path.join(OUTPUT_DIR, "C1_Gerer_patrimoine_informatique.docx")
    doc.save(path)
    print(f"[OK] {path}")


# ─── Fiche 2 : Répondre aux incidents et demandes ───────────────────────────

def fiche_2():
    doc = create_doc()
    add_header_info(doc)

    add_title(doc, "Réalisation professionnelle : Liaison Windows Server avec plusieurs clients Windows et gestion des tickets GLPI")

    add_section(doc, "Compétences mises en œuvre")
    add_body(doc,
        "• Collecter, suivre et orienter des demandes\n"
        "• Traiter des demandes concernant les services réseau et système, applicatifs\n"
        "• Traiter des demandes concernant les applications"
    )

    add_section(doc, "Contexte")
    add_body(doc,
        "Dans le cadre de ma formation, j'ai configuré une infrastructure client-serveur sous Windows "
        "en reliant un contrôleur de domaine Windows Server 2019 à plusieurs postes clients Windows 10/11. "
        "Parallèlement, j'ai utilisé le module Helpdesk de GLPI pour simuler la prise en charge "
        "et le suivi des incidents signalés par les utilisateurs."
    )

    add_section(doc, "Démarche suivie")
    add_body(doc,
        "1. Installation et configuration de Windows Server 2019 :\n"
        "   - Promotion en contrôleur de domaine Active Directory (AD DS)\n"
        "   - Configuration du service DNS intégré à AD\n"
        "   - Création des utilisateurs et groupes de sécurité"
    )
    add_code_block(doc,
        "# PowerShell — Promotion en contrôleur de domaine\n"
        "Install-WindowsFeature AD-Domain-Services -IncludeManagementTools\n"
        "Install-ADDSForest -DomainName 'solal.local' -InstallDns"
    )
    add_body(doc,
        "2. Jonction des postes clients Windows au domaine :\n"
        "   - Paramétrage du DNS sur les clients (pointage vers le serveur AD)\n"
        "   - Jonction au domaine via Paramètres > Système > Domaine\n"
        "   - Application des stratégies de groupe (GPO) pour la configuration centralisée"
    )
    add_body(doc,
        "3. Prise en main du helpdesk GLPI :\n"
        "   - Création de catégories de tickets (incident réseau, problème applicatif, demande matériel)\n"
        "   - Simulation d'ouverture de tickets par des utilisateurs fictifs\n"
        "   - Attribution des tickets aux techniciens, suivi des SLA et clôture après résolution"
    )

    add_section(doc, "Conclusion")
    add_body(doc,
        "Cette réalisation m'a permis de comprendre l'architecture d'un environnement Active Directory "
        "et la gestion centralisée des comptes utilisateurs. La pratique du helpdesk GLPI m'a sensibilisé "
        "aux processus ITIL de gestion des incidents et des demandes. J'ai développé ma rigueur dans "
        "le suivi documentaire des interventions, compétence essentielle en milieu professionnel."
    )

    path = os.path.join(OUTPUT_DIR, "C2_Repondre_incidents_demandes.docx")
    doc.save(path)
    print(f"[OK] {path}")


# ─── Fiche 3 : Développer la présence en ligne ──────────────────────────────

def fiche_3():
    doc = create_doc()
    add_header_info(doc)

    add_title(doc, "Réalisation professionnelle : Création d'un site internet en équipe et déploiement sur Azure")

    add_section(doc, "Compétences mises en œuvre")
    add_body(doc,
        "• Participer à la valorisation de l'image de l'organisation sur les médias numériques "
        "en tenant compte du cadre juridique et des enjeux économiques\n"
        "• Référencer les services en ligne de l'organisation et mesurer leur visibilité\n"
        "• Participer à l'évolution d'un site Web exploitant les données de l'organisation"
    )

    add_section(doc, "Contexte")
    add_body(doc,
        "Dans le cadre d'un projet collectif réalisé en formation à l'EFREI, notre équipe a développé "
        "un site web vitrine pour une organisation fictive. L'objectif était de mettre en ligne un service "
        "numérique représentatif, accessible depuis Internet, et d'en mesurer la visibilité. "
        "J'ai ensuite déployé un serveur web sur Microsoft Azure pour héberger le site."
    )

    add_section(doc, "Démarche suivie")
    add_body(doc,
        "1. Conception du site web en équipe :\n"
        "   - Définition de la charte graphique et de l'arborescence\n"
        "   - Développement des pages en HTML5, CSS3 et JavaScript\n"
        "   - Intégration d'un formulaire de contact et d'une page de présentation des services"
    )
    add_code_block(doc,
        "<!-- Structure HTML de base -->\n"
        "<!DOCTYPE html>\n"
        "<html lang='fr'>\n"
        "  <head><meta charset='UTF-8'><title>Organisation</title></head>\n"
        "  <body>...</body>\n"
        "</html>"
    )
    add_body(doc,
        "2. Déploiement sur Microsoft Azure :\n"
        "   - Création d'une machine virtuelle Ubuntu dans la région France Centre\n"
        "   - Installation et configuration d'Apache2 comme serveur web\n"
        "   - Configuration des règles NSG (Network Security Group) pour ouvrir les ports 80 et 443"
    )
    add_code_block(doc,
        "# Installation Apache sur la VM Azure\n"
        "sudo apt update\n"
        "sudo apt install apache2 -y\n"
        "sudo systemctl enable apache2\n"
        "# Copie des fichiers du site\n"
        "sudo cp -r /home/solal/site/* /var/www/html/"
    )
    add_body(doc,
        "3. Référencement et mesure de la visibilité :\n"
        "   - Vérification de l'accessibilité via l'adresse IP publique Azure\n"
        "   - Analyse des en-têtes HTTP et du SEO de base (balises meta, title, description)\n"
        "   - Respect du RGPD : mentions légales, politique de confidentialité, gestion des cookies"
    )

    add_section(doc, "Conclusion")
    add_body(doc,
        "Ce projet m'a permis de maîtriser le déploiement d'un service web de bout en bout : "
        "du développement front-end jusqu'à l'hébergement cloud sur Azure. "
        "J'ai pris conscience des enjeux juridiques liés à la publication d'un site (RGPD, droits d'auteur) "
        "et des bonnes pratiques de référencement naturel. Le travail en équipe a renforcé mes capacités "
        "de collaboration et de communication technique."
    )

    path = os.path.join(OUTPUT_DIR, "C3_Developper_presence_en_ligne.docx")
    doc.save(path)
    print(f"[OK] {path}")


# ─── Fiche 4 : Travailler en mode projet ────────────────────────────────────

def fiche_4():
    doc = create_doc()
    add_header_info(doc)

    add_title(doc, "Réalisation professionnelle : Gestion de projet avec Trello et audit réseau avec Nmap")

    add_section(doc, "Compétences mises en œuvre")
    add_body(doc,
        "• Analyser les objectifs et les modalités d'organisation d'un projet\n"
        "• Planifier les activités\n"
        "• Évaluer les indicateurs de suivi d'un projet et analyser les écarts"
    )

    add_section(doc, "Contexte")
    add_body(doc,
        "Lors de plusieurs projets réalisés en formation, j'ai utilisé Trello comme outil de gestion "
        "de projet en mode Kanban pour coordonner le travail en équipe. En parallèle, j'ai mené un audit "
        "réseau avec Nmap dans le cadre d'une analyse de l'infrastructure d'une organisation fictive, "
        "nécessitant une planification rigoureuse des étapes d'investigation."
    )

    add_section(doc, "Démarche suivie")
    add_body(doc,
        "1. Analyse des objectifs du projet :\n"
        "   - Réunion de cadrage pour définir le périmètre, les livrables et les délais\n"
        "   - Identification des parties prenantes et de leurs attentes\n"
        "   - Découpage du projet en tâches élémentaires (WBS)"
    )
    add_body(doc,
        "2. Planification via Trello :\n"
        "   - Création d'un tableau Kanban avec les colonnes : À faire / En cours / En révision / Terminé\n"
        "   - Attribution des tâches aux membres de l'équipe avec dates d'échéance\n"
        "   - Mise en place d'étiquettes de priorité (haute, moyenne, basse)"
    )
    add_body(doc,
        "3. Audit réseau avec Nmap :\n"
        "   - Définition du périmètre de scan et des autorisations nécessaires\n"
        "   - Lancement du scan de découverte des hôtes actifs sur le réseau"
    )
    add_code_block(doc,
        "# Découverte des hôtes actifs\n"
        "nmap -sn 192.168.1.0/24\n\n"
        "# Scan complet des services et versions\n"
        "nmap -sV -sC -O 192.168.1.0/24 -oN rapport_audit.txt\n\n"
        "# Détection des ports ouverts sur une cible spécifique\n"
        "nmap -p 1-65535 -T4 192.168.1.50"
    )
    add_body(doc,
        "4. Suivi et analyse des écarts :\n"
        "   - Points d'avancement hebdomadaires en équipe\n"
        "   - Analyse des retards et re-planification des tâches bloquées\n"
        "   - Rédaction d'un compte-rendu de fin de sprint"
    )

    add_section(doc, "Conclusion")
    add_body(doc,
        "La gestion de projet via Trello m'a appris à structurer mon travail, à prioriser les tâches "
        "et à communiquer efficacement au sein d'une équipe. La méthode Kanban s'est révélée adaptée "
        "aux projets informatiques de courte durée. L'audit Nmap a illustré l'importance d'une approche "
        "méthodique et documentée : chaque étape doit être planifiée, exécutée et analysée pour en tirer "
        "des conclusions exploitables."
    )

    path = os.path.join(OUTPUT_DIR, "C4_Travailler_en_mode_projet.docx")
    doc.save(path)
    print(f"[OK] {path}")


# ─── Fiche 5 : Mettre à disposition un service informatique ─────────────────

def fiche_5():
    doc = create_doc()
    add_header_info(doc)

    add_title(doc, "Réalisation professionnelle : Déploiement de GLPI et d'un serveur web sur Microsoft Azure")

    add_section(doc, "Compétences mises en œuvre")
    add_body(doc,
        "• Réaliser les tests d'intégration et d'acceptation d'un service\n"
        "• Déployer un service\n"
        "• Accompagner les utilisateurs dans la mise en place d'un service"
    )

    add_section(doc, "Contexte")
    add_body(doc,
        "Dans le cadre de ma formation BTS SIO, j'ai été chargé de déployer deux services informatiques "
        "accessibles depuis Internet : une instance GLPI pour la gestion de parc et un serveur web "
        "vitrine, tous deux hébergés sur Microsoft Azure. L'objectif était de rendre ces services "
        "disponibles aux utilisateurs après validation complète."
    )

    add_section(doc, "Démarche suivie")
    add_body(doc,
        "1. Provisionnement de l'infrastructure Azure :\n"
        "   - Création d'un groupe de ressources et d'un réseau virtuel (VNet)\n"
        "   - Déploiement d'une VM Ubuntu Server 22.04 LTS\n"
        "   - Configuration des NSG : ouverture des ports 22 (SSH), 80 (HTTP), 443 (HTTPS)"
    )
    add_code_block(doc,
        "# Connexion SSH à la VM Azure\n"
        "ssh -i ~/.ssh/azure_key.pem solal@<IP_PUBLIQUE>\n\n"
        "# Installation de la stack LAMP\n"
        "sudo apt update && sudo apt install -y apache2 mariadb-server php php-mysql\n"
        "sudo systemctl enable --now apache2 mariadb"
    )
    add_body(doc,
        "2. Installation et configuration de GLPI 10.x sur la VM :\n"
        "   - Création de la base de données dédiée\n"
        "   - Configuration du virtualhost Apache pour GLPI\n"
        "   - Finalisation via l'assistant d'installation web"
    )
    add_code_block(doc,
        "mysql -u root -p -e \"CREATE DATABASE glpidb; \\\n"
        "  GRANT ALL ON glpidb.* TO 'glpiuser'@'localhost' IDENTIFIED BY 'motdepasse';\""
    )
    add_body(doc,
        "3. Tests d'intégration et d'acceptation :\n"
        "   - Vérification de l'accès depuis différents navigateurs et réseaux\n"
        "   - Test de création d'inventaire et de ticket depuis un compte utilisateur\n"
        "   - Contrôle des performances et des logs Apache\n"
        "   - Validation avec les utilisateurs fictifs : recueil de retours et corrections"
    )
    add_body(doc,
        "4. Accompagnement des utilisateurs :\n"
        "   - Rédaction d'une documentation utilisateur simplifiée\n"
        "   - Présentation des fonctionnalités principales (inventaire, tickets, rapports)\n"
        "   - Formation courte aux bonnes pratiques d'utilisation"
    )

    add_section(doc, "Conclusion")
    add_body(doc,
        "Ce projet m'a permis d'appréhender le cycle complet de mise à disposition d'un service : "
        "de l'infrastructure cloud jusqu'à la formation des utilisateurs. L'hébergement sur Azure "
        "m'a familiarisé avec les concepts de cloud public (IaaS), la gestion des ressources et la "
        "sécurisation des accès. Les tests d'acceptation ont montré l'importance de valider un service "
        "du point de vue utilisateur avant toute mise en production."
    )

    path = os.path.join(OUTPUT_DIR, "C5_Mettre_a_disposition_service.docx")
    doc.save(path)
    print(f"[OK] {path}")


# ─── Fiche 6 : Organiser son développement professionnel ────────────────────

def fiche_6():
    doc = create_doc()
    add_header_info(doc)

    add_title(doc, "Réalisation professionnelle : Création du portfolio numérique et veille technologique")

    add_section(doc, "Compétences mises en œuvre")
    add_body(doc,
        "• Mettre en place son environnement d'apprentissage personnel\n"
        "• Mettre en œuvre des outils et stratégies de veille informationnelle\n"
        "• Gérer son identité professionnelle\n"
        "• Développer son projet professionnel"
    )

    add_section(doc, "Contexte")
    add_body(doc,
        "Dans le cadre du BTS SIO, il m'a été demandé de constituer un portfolio numérique recensant "
        "l'ensemble de mes réalisations professionnelles et compétences acquises. J'ai développé ce "
        "portfolio sous forme de site web moderne (HTML/CSS/JS), hébergeable en ligne, et j'ai mis en "
        "place une stratégie de veille informationnelle pour suivre l'évolution des technologies "
        "liées à mon domaine (cybersécurité, cloud, systèmes et réseaux)."
    )

    add_section(doc, "Démarche suivie")
    add_body(doc,
        "1. Construction du portfolio numérique :\n"
        "   - Conception de la structure : page d'accueil, compétences BTS, projets, contact\n"
        "   - Développement en HTML5 / CSS3 (animations, design responsive) et JavaScript\n"
        "   - Intégration du tableau officiel des compétences BTS SIO (PDF embarqué)\n"
        "   - Mise en valeur des projets : GLPI, Azure, Nmap, Trello, Windows Server"
    )
    add_code_block(doc,
        "<!-- Exemple d'intégration d'une compétence dans le portfolio -->\n"
        "<div class='comp-card'>\n"
        "  <div class='comp-badge'>C1</div>\n"
        "  <h3>Recensement des ressources numériques</h3>\n"
        "  <div class='comp-tools'>\n"
        "    <span>GLPI</span><span>Nmap</span><span>Debian</span>\n"
        "  </div>\n"
        "</div>"
    )
    add_body(doc,
        "2. Mise en place de la veille informationnelle :\n"
        "   - Abonnement à des flux RSS (The Hacker News, LeMagIT, CERT-FR)\n"
        "   - Suivi de communautés sur GitHub et forums spécialisés (Reddit r/netsec, Stack Overflow)\n"
        "   - Utilisation de Feedly pour centraliser les sources\n"
        "   - Lecture régulière des bulletins de sécurité ANSSI"
    )
    add_body(doc,
        "3. Gestion de l'identité professionnelle :\n"
        "   - Mise à jour du profil LinkedIn avec les compétences et certifications acquises\n"
        "   - Référencement du portfolio en ligne comme vitrine professionnelle\n"
        "   - Définition d'un projet professionnel orienté vers la cybersécurité et l'administration systèmes"
    )

    add_section(doc, "Conclusion")
    add_body(doc,
        "La création du portfolio m'a permis de prendre du recul sur l'ensemble de mes réalisations "
        "et de les valoriser de manière professionnelle. Cet exercice a renforcé mes compétences en "
        "développement web tout en structurant mon projet professionnel autour de la cybersécurité "
        "et de l'administration des systèmes. La veille informationnelle mise en place me permet de "
        "rester informé des évolutions technologiques et des nouvelles menaces, compétence indispensable "
        "dans le secteur de l'informatique."
    )

    path = os.path.join(OUTPUT_DIR, "C6_Organiser_developpement_professionnel.docx")
    doc.save(path)
    print(f"[OK] {path}")


# ─── Main ────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Génération des fiches de compétences BTS SIO...\n")
    fiche_1()
    fiche_2()
    fiche_3()
    fiche_4()
    fiche_5()
    fiche_6()
    print(f"\nTous les documents ont été générés dans : {OUTPUT_DIR}")
