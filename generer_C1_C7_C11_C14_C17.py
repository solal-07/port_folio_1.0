"""
Génération des fiches de compétences BTS SIO — Solal BEN ZAQUI
Compétences : C1, C7, C11, C14, C17
Format conforme à l'exemple de documentation fourni.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "fiches_competences")
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ─── Helpers ────────────────────────────────────────────────────────────────

def add_title(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.size = Pt(22)
    p.runs[0].font.color.rgb = RGBColor(31, 78, 121)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_section(doc, title):
    p = doc.add_heading(title, level=2)
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.color.rgb = RGBColor(68, 114, 196)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.runs[0].font.size = Pt(11)
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(40, 40, 40)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def create_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    return doc


def add_header_info(doc, nom_candidat="BEN ZAQUI SOLAL", centre="BTS SIO — EFREI"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Candidat : {nom_candidat}  |  {centre}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)
    doc.add_paragraph()


# ─── C1 : Recenser et identifier les ressources numériques ──────────────────

def fiche_C1():
    doc = create_doc()
    add_header_info(doc)

    add_title(doc, "Réalisation professionnelle : Recensement du parc informatique avec GLPI sur VM Debian 12")

    add_section(doc, "Compétence mise en œuvre")
    add_body(doc,
        "C1 — Recenser et identifier les ressources numériques\n"
        "(Bloc 1 : Gérer le patrimoine informatique)"
    )

    add_section(doc, "Contexte")
    add_body(doc,
        "Dans le cadre de ma formation BTS SIO option SISR à l'EFREI, j'ai été amené à mettre en place "
        "un système de gestion de parc informatique (ITSM) afin d'inventorier et de suivre l'ensemble "
        "des ressources numériques d'une organisation fictive. L'objectif était d'améliorer la "
        "traçabilité des équipements, des logiciels et des licences, en centralisant toutes les "
        "informations dans un outil unique et accessible."
    )

    add_section(doc, "Démarche suivie")
    add_body(doc,
        "1. Création et configuration d'une machine virtuelle sous Debian 12 (VirtualBox) :\n"
        "   - Allocation des ressources (RAM, CPU, stockage)\n"
        "   - Configuration du réseau en mode pont pour l'accessibilité depuis le LAN\n"
        "   - Mise à jour du système"
    )
    add_code_block(doc,
        "sudo apt update && sudo apt upgrade -y"
    )
    add_body(doc,
        "2. Installation de la pile LAMP (Apache2, MariaDB, PHP) nécessaire au fonctionnement de GLPI :"
    )
    add_code_block(doc,
        "sudo apt install apache2 mariadb-server php php-mysql php-curl php-gd \\\n"
        "                 php-intl php-ldap php-mbstring php-xml php-zip \\\n"
        "                 libapache2-mod-php -y\n"
        "sudo systemctl enable --now apache2 mariadb"
    )
    add_body(doc,
        "3. Téléchargement, installation et configuration de GLPI 10.x :"
    )
    add_code_block(doc,
        "wget https://github.com/glpi-project/glpi/releases/download/10.0.12/glpi-10.0.12.tgz\n"
        "tar -xzf glpi-10.0.12.tgz -C /var/www/html/\n"
        "sudo chown -R www-data:www-data /var/www/html/glpi\n"
        "# Création de la base de données\n"
        "mysql -u root -p -e \"CREATE DATABASE glpidb CHARACTER SET utf8mb4;\n"
        "  GRANT ALL ON glpidb.* TO 'glpiuser'@'localhost' IDENTIFIED BY 'Motdepasse123!';\""
    )
    add_body(doc,
        "4. Déploiement de l'agent GLPI sur les machines clientes pour la remontée automatique "
        "des informations matérielles et logicielles :"
    )
    add_code_block(doc,
        "sudo apt install glpi-agent\n"
        "# Édition du fichier de configuration\n"
        "sudo nano /etc/glpi-agent/agent.cfg\n"
        "  server = http://192.168.1.100/glpi\n"
        "sudo systemctl restart glpi-agent"
    )
    add_body(doc,
        "5. Recensement et identification des ressources dans l'interface GLPI :\n"
        "   - Création des catégories d'équipements (ordinateurs, imprimantes, switchs, serveurs)\n"
        "   - Consultation de l'inventaire automatiquement remonté par les agents\n"
        "   - Vérification et enrichissement des fiches : numéros de série, adresses IP, OS, logiciels\n"
        "   - Attribution des équipements aux utilisateurs et aux lieux (salles, services)"
    )

    add_section(doc, "Conclusion")
    add_body(doc,
        "Cette réalisation m'a permis de maîtriser le déploiement complet de GLPI sur un environnement "
        "Linux Debian et de comprendre l'importance du recensement du patrimoine numérique dans une "
        "organisation. Grâce à l'agent GLPI, l'inventaire se met à jour automatiquement, ce qui garantit "
        "une vision fiable et en temps réel du parc. J'ai également développé ma capacité à lire et "
        "interpréter les informations remontées pour en tirer des indicateurs utiles à la gestion du SI."
    )

    path = os.path.join(OUTPUT_DIR, "C1_Recenser_ressources_numeriques.docx")
    doc.save(path)
    print(f"[OK] {path}")


# ─── C7 : Collecter, suivre et orienter des demandes ────────────────────────

def fiche_C7():
    doc = create_doc()
    add_header_info(doc)

    add_title(doc, "Réalisation professionnelle : Gestion des tickets et des demandes avec le helpdesk GLPI")

    add_section(doc, "Compétence mise en œuvre")
    add_body(doc,
        "C7 — Collecter, suivre et orienter des demandes\n"
        "(Bloc 2 : Répondre aux incidents et aux demandes d'assistance et d'évolution)"
    )

    add_section(doc, "Contexte")
    add_body(doc,
        "Dans le cadre de ma formation BTS SIO option SISR, j'ai utilisé le module Helpdesk de GLPI "
        "pour simuler la prise en charge complète des demandes et incidents signalés par les utilisateurs "
        "d'une organisation fictive. L'objectif était de s'approprier un processus ITIL de gestion des "
        "tickets : de la réception de la demande jusqu'à sa clôture, en passant par l'orientation "
        "vers le bon intervenant et le suivi de la résolution."
    )

    add_section(doc, "Démarche suivie")
    add_body(doc,
        "1. Configuration du helpdesk dans GLPI :\n"
        "   - Création des catégories de tickets (incident réseau, problème applicatif, demande matériel, accès)\n"
        "   - Paramétrage des SLA (Service Level Agreement) : délais de prise en charge et de résolution\n"
        "   - Création des groupes de techniciens et des règles d'attribution automatique"
    )
    add_body(doc,
        "2. Collecte des demandes :\n"
        "   - Simulation d'ouverture de tickets par des utilisateurs fictifs via l'interface self-service\n"
        "   - Réception et qualification des demandes : type (incident / demande), urgence, impact\n"
        "   - Vérification des informations manquantes et contact de l'utilisateur si nécessaire"
    )
    add_body(doc,
        "3. Suivi et orientation des demandes :\n"
        "   - Attribution des tickets aux techniciens compétents selon la catégorie\n"
        "   - Mise à jour du statut en temps réel : Nouveau → En cours → En attente → Résolu → Clôturé\n"
        "   - Ajout de notes internes pour tracer les actions réalisées"
    )
    add_code_block(doc,
        "# Exemple de flux de traitement d'un ticket GLPI\n"
        "Ticket #42 — Panne réseau poste de travail\n"
        "  Statut initial  : Nouveau\n"
        "  Priorité        : Haute (impact utilisateur bloqué)\n"
        "  Assigné à       : Groupe Réseau\n"
        "  Action effectuée: Vérification câblage + renouvellement bail DHCP\n"
        "  Statut final    : Résolu — clôturé après validation utilisateur"
    )
    add_body(doc,
        "4. Analyse et bilan des demandes :\n"
        "   - Consultation des tableaux de bord GLPI (nombre de tickets par catégorie, délai moyen)\n"
        "   - Identification des incidents récurrents pour proposer des actions préventives\n"
        "   - Rédaction d'un compte-rendu hebdomadaire des interventions"
    )

    add_section(doc, "Conclusion")
    add_body(doc,
        "Cette réalisation m'a sensibilisé aux bonnes pratiques ITIL de gestion des incidents et des "
        "demandes. J'ai appris à qualifier, prioriser et orienter chaque ticket vers le bon interlocuteur, "
        "tout en maintenant une communication claire avec l'utilisateur. Le suivi rigoureux des tickets "
        "dans GLPI garantit la traçabilité des interventions et améliore la qualité de service rendu "
        "par le support informatique."
    )

    path = os.path.join(OUTPUT_DIR, "C7_Collecter_suivre_orienter_demandes.docx")
    doc.save(path)
    print(f"[OK] {path}")


# ─── C11 : Référencer les services en ligne et mesurer leur visibilité ───────

def fiche_C11():
    doc = create_doc()
    add_header_info(doc)

    add_title(doc, "Réalisation professionnelle : Référencement et mesure de la visibilité d'un site internet déployé sur Azure")

    add_section(doc, "Compétence mise en œuvre")
    add_body(doc,
        "C11 — Référencer les services en ligne de l'organisation et mesurer leur visibilité\n"
        "(Bloc 3 : Développer la présence en ligne de l'organisation)"
    )

    add_section(doc, "Contexte")
    add_body(doc,
        "Dans le cadre d'un projet collectif réalisé à l'EFREI, notre équipe a développé un site web "
        "vitrine pour une organisation fictive et l'a déployé sur Microsoft Azure. Une fois le site "
        "en ligne, j'ai été chargé d'analyser son référencement naturel (SEO), de mesurer sa visibilité "
        "sur le web et de proposer des améliorations pour optimiser son positionnement et son accessibilité."
    )

    add_section(doc, "Démarche suivie")
    add_body(doc,
        "1. Déploiement du site sur Microsoft Azure :\n"
        "   - Création d'une machine virtuelle Ubuntu dans la région France Centre\n"
        "   - Installation et configuration d'Apache2 comme serveur web\n"
        "   - Ouverture des ports HTTP (80) et HTTPS (443) via les règles NSG"
    )
    add_code_block(doc,
        "sudo apt update && sudo apt install apache2 -y\n"
        "sudo systemctl enable --now apache2\n"
        "# Copie des fichiers du site dans le répertoire web\n"
        "sudo cp -r /home/solal/site/* /var/www/html/"
    )
    add_body(doc,
        "2. Référencement du service en ligne :\n"
        "   - Vérification de l'accessibilité via l'adresse IP publique Azure et le nom DNS\n"
        "   - Optimisation des balises HTML essentielles au SEO :"
    )
    add_code_block(doc,
        "<!-- Balises SEO dans le <head> -->\n"
        "<meta name='description' content='Organisation fictive — Services informatiques'>\n"
        "<meta name='keywords' content='informatique, services, BTS SIO, réseau'>\n"
        "<meta name='author' content='BEN ZAQUI Solal'>\n"
        "<title>Organisation Fictive | Accueil</title>\n"
        "<!-- Open Graph pour les réseaux sociaux -->\n"
        "<meta property='og:title' content='Organisation Fictive'>\n"
        "<meta property='og:description' content='Votre prestataire informatique'>"
    )
    add_body(doc,
        "3. Mesure de la visibilité :\n"
        "   - Analyse des en-têtes HTTP avec l'outil curl pour vérifier les codes de réponse\n"
        "   - Test de performance et de temps de chargement avec les outils de développement du navigateur\n"
        "   - Vérification de la conformité aux bonnes pratiques SEO (structure des titres H1/H2, "
        "attributs alt des images, sitemap.xml, robots.txt)\n"
        "   - Contrôle de l'accessibilité du site depuis différents réseaux (4G, Wi-Fi, VPN)"
    )
    add_code_block(doc,
        "# Vérification du code de réponse HTTP\n"
        "curl -I http://<IP_PUBLIQUE_AZURE>\n"
        "# Résultat attendu : HTTP/1.1 200 OK\n\n"
        "# Test de résolution DNS\n"
        "nslookup monsite.azure.com"
    )
    add_body(doc,
        "4. Respect du cadre juridique :\n"
        "   - Ajout des mentions légales obligatoires (éditeur, hébergeur)\n"
        "   - Mise en place d'une politique de confidentialité conforme au RGPD\n"
        "   - Intégration d'un bandeau de consentement aux cookies"
    )

    add_section(doc, "Conclusion")
    add_body(doc,
        "Cette réalisation m'a permis de comprendre les mécanismes du référencement naturel et les "
        "facteurs qui influencent la visibilité d'un service en ligne. J'ai appris à analyser "
        "techniquement un site web du point de vue du SEO et à mettre en œuvre les bonnes pratiques "
        "pour améliorer son accessibilité. J'ai également pris conscience des obligations légales "
        "liées à la publication de contenu en ligne (RGPD, mentions légales)."
    )

    path = os.path.join(OUTPUT_DIR, "C11_Referencer_services_en_ligne.docx")
    doc.save(path)
    print(f"[OK] {path}")


# ─── C14 : Planifier les activités ──────────────────────────────────────────

def fiche_C14():
    doc = create_doc()
    add_header_info(doc)

    add_title(doc, "Réalisation professionnelle : Planification des activités de projet avec Trello")

    add_section(doc, "Compétence mise en œuvre")
    add_body(doc,
        "C14 — Planifier les activités\n"
        "(Bloc 4 : Travailler en mode projet)"
    )

    add_section(doc, "Contexte")
    add_body(doc,
        "Lors de plusieurs projets réalisés en formation à l'EFREI (création du site internet en équipe, "
        "déploiement de GLPI, audit réseau avec Nmap), j'ai utilisé Trello comme outil de gestion de "
        "projet en mode Kanban. L'objectif était de planifier les activités de manière structurée, "
        "de répartir les tâches entre les membres de l'équipe et de suivre l'avancement en temps réel "
        "pour respecter les délais fixés."
    )

    add_section(doc, "Démarche suivie")
    add_body(doc,
        "1. Analyse des objectifs et découpage du projet :\n"
        "   - Réunion de cadrage pour définir le périmètre, les livrables et les contraintes\n"
        "   - Identification des grandes phases du projet (analyse, développement, test, déploiement)\n"
        "   - Décomposition en tâches élémentaires (Work Breakdown Structure — WBS)\n"
        "   - Estimation de la durée et de la charge de chaque tâche"
    )
    add_body(doc,
        "2. Création et paramétrage du tableau Trello :\n"
        "   - Mise en place des colonnes Kanban : À faire / En cours / En révision / Terminé\n"
        "   - Création des cartes pour chaque tâche avec une description détaillée\n"
        "   - Attribution des cartes aux membres de l'équipe\n"
        "   - Définition des dates d'échéance et des étiquettes de priorité (rouge = urgent, jaune = normal, vert = faible)"
    )
    add_code_block(doc,
        "Exemple d'organisation du tableau Trello — Projet Site Web\n"
        "┌──────────────┬──────────────┬──────────────┬──────────────┐\n"
        "│   À faire    │  En cours    │ En révision  │   Terminé    │\n"
        "├──────────────┼──────────────┼──────────────┼──────────────┤\n"
        "│ Page contact │ Page accueil │ Page about   │ Maquette     │\n"
        "│ Tests finaux │ CSS mobile   │              │ Charte graph.│\n"
        "│ Déploiement  │              │              │ Structure HTML│\n"
        "└──────────────┴──────────────┴──────────────┴──────────────┘"
    )
    add_body(doc,
        "3. Suivi de l'avancement et gestion des écarts :\n"
        "   - Points d'avancement quotidiens (stand-up) de 15 minutes en équipe\n"
        "   - Déplacement des cartes au fil de l'avancement des tâches\n"
        "   - Identification des blocages et re-planification des tâches en retard\n"
        "   - Mise à jour des dates d'échéance si nécessaire et communication à l'équipe"
    )
    add_body(doc,
        "4. Clôture et bilan du projet :\n"
        "   - Vérification que toutes les tâches sont dans la colonne « Terminé »\n"
        "   - Rétrospective : ce qui a bien fonctionné, ce qui peut être amélioré\n"
        "   - Archivage du tableau Trello comme référence documentaire"
    )

    add_section(doc, "Conclusion")
    add_body(doc,
        "L'utilisation de Trello comme outil de planification m'a appris à structurer un projet "
        "en tâches concrètes et à les répartir efficacement au sein d'une équipe. La méthode Kanban "
        "offre une visibilité immédiate sur l'état du projet et facilite la détection précoce des "
        "retards. J'ai développé ma rigueur dans la planification et ma capacité à adapter le "
        "planning en cours de projet sans perdre de vue les objectifs finaux."
    )

    path = os.path.join(OUTPUT_DIR, "C14_Planifier_les_activites.docx")
    doc.save(path)
    print(f"[OK] {path}")


# ─── C17 : Déployer un service ───────────────────────────────────────────────

def fiche_C17():
    doc = create_doc()
    add_header_info(doc)

    add_title(doc, "Réalisation professionnelle : Déploiement de services informatiques sur Microsoft Azure")

    add_section(doc, "Compétence mise en œuvre")
    add_body(doc,
        "C17 — Déployer un service\n"
        "(Bloc 5 : Mettre à disposition des utilisateurs un service informatique)"
    )

    add_section(doc, "Contexte")
    add_body(doc,
        "Dans le cadre de ma formation BTS SIO option SISR à l'EFREI, j'ai été chargé de déployer "
        "deux services informatiques accessibles depuis Internet : une instance GLPI pour la gestion "
        "de parc informatique et un serveur web vitrine, tous deux hébergés sur Microsoft Azure. "
        "L'objectif était de rendre ces services pleinement opérationnels et accessibles aux "
        "utilisateurs, en suivant une procédure de déploiement rigoureuse."
    )

    add_section(doc, "Démarche suivie")
    add_body(doc,
        "1. Provisionnement de l'infrastructure sur Microsoft Azure :\n"
        "   - Connexion au portail Azure et création d'un groupe de ressources dédié\n"
        "   - Déploiement d'une machine virtuelle Ubuntu Server 22.04 LTS (taille B2s)\n"
        "   - Configuration du réseau virtuel (VNet) et des règles de sécurité (NSG) :\n"
        "     ouverture des ports 22 (SSH), 80 (HTTP), 443 (HTTPS)\n"
        "   - Récupération de l'adresse IP publique"
    )
    add_code_block(doc,
        "# Connexion SSH sécurisée à la VM Azure\n"
        "ssh -i ~/.ssh/azure_key.pem azureuser@<IP_PUBLIQUE>\n\n"
        "# Mise à jour du système\n"
        "sudo apt update && sudo apt upgrade -y"
    )
    add_body(doc,
        "2. Déploiement de la pile LAMP (prérequis pour GLPI et le site web) :"
    )
    add_code_block(doc,
        "sudo apt install -y apache2 mariadb-server php php-mysql php-curl \\\n"
        "                    php-gd php-intl php-mbstring php-xml php-zip \\\n"
        "                    libapache2-mod-php\n"
        "sudo systemctl enable --now apache2 mariadb"
    )
    add_body(doc,
        "3. Déploiement de GLPI 10.x :\n"
        "   - Téléchargement de l'archive et extraction dans le répertoire web\n"
        "   - Création de la base de données et de l'utilisateur MariaDB dédié\n"
        "   - Configuration du virtualhost Apache et finalisation via l'assistant web"
    )
    add_code_block(doc,
        "wget https://github.com/glpi-project/glpi/releases/download/10.0.12/glpi-10.0.12.tgz\n"
        "tar -xzf glpi-10.0.12.tgz -C /var/www/html/\n"
        "sudo chown -R www-data:www-data /var/www/html/glpi\n\n"
        "mysql -u root -p -e \"\n"
        "  CREATE DATABASE glpidb CHARACTER SET utf8mb4;\n"
        "  GRANT ALL ON glpidb.* TO 'glpiuser'@'localhost' IDENTIFIED BY 'Motdepasse123!';\n"
        "\""
    )
    add_body(doc,
        "4. Déploiement du site web vitrine :\n"
        "   - Transfert des fichiers HTML/CSS/JS via SCP depuis le poste local\n"
        "   - Configuration d'un virtualhost Apache dédié\n"
        "   - Vérification de l'accessibilité depuis Internet"
    )
    add_code_block(doc,
        "# Transfert des fichiers du site\n"
        "scp -i ~/.ssh/azure_key.pem -r ./site/* azureuser@<IP_PUBLIQUE>:/var/www/html/\n\n"
        "# Vérification du service Apache\n"
        "sudo systemctl status apache2\n"
        "curl -I http://<IP_PUBLIQUE>   # doit retourner HTTP/1.1 200 OK"
    )
    add_body(doc,
        "5. Tests de déploiement et validation :\n"
        "   - Accès aux deux services depuis différents navigateurs et réseaux\n"
        "   - Vérification des fonctionnalités principales de GLPI (inventaire, tickets)\n"
        "   - Contrôle des logs Apache pour détecter d'éventuelles erreurs\n"
        "   - Validation du bon fonctionnement avec des utilisateurs tests"
    )

    add_section(doc, "Conclusion")
    add_body(doc,
        "Ce projet m'a permis d'appréhender le cycle complet de déploiement d'un service informatique "
        "sur une infrastructure cloud. L'utilisation de Microsoft Azure m'a familiarisé avec les "
        "concepts IaaS (Infrastructure as a Service) : provisionnement de VM, gestion des ressources "
        "réseau et sécurisation des accès. La rigueur de la procédure de déploiement — de "
        "l'installation à la validation — est essentielle pour garantir un service fiable et "
        "disponible pour les utilisateurs finaux."
    )

    path = os.path.join(OUTPUT_DIR, "C17_Deployer_un_service.docx")
    doc.save(path)
    print(f"[OK] {path}")


# ─── Main ────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Génération des fiches de compétences BTS SIO (C1, C7, C11, C14, C17)...\n")
    fiche_C1()
    fiche_C7()
    fiche_C11()
    fiche_C14()
    fiche_C17()
    print(f"\nTous les documents ont été générés dans : {OUTPUT_DIR}")
